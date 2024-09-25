# Import required libraries
import speech_recognition as sr
from tkinter import *
from tkinter import ttk, messagebox, filedialog
import pyperclip
from googletrans import Translator as GoogleTranslator, LANGUAGES
from deep_translator import (
    GoogleTranslator as DeepGoogleTranslator,
    MicrosoftTranslator,
    PonsTranslator,
    LingueeTranslator,
    MyMemoryTranslator
)
from docx import Document
import PyPDF2
import json
from gtts import gTTS
import os
from langdetect import detect, DetectorFactory

# Set seed for consistent language detection
DetectorFactory.seed = 0

# Global variables to store translation history, custom languages, and Listbox for displaying history
translation_history = []
custom_languages = {}
history_listbox = None


def change(text, src="english", dest="marathi", translator_choice="Google (deep-translator)"):
    """Handles the translation process based on the selected translator."""
    try:
        if translator_choice == "Google (deep-translator)":
            translated = DeepGoogleTranslator(source=src, target=dest).translate(text)
        elif translator_choice == "Google (googletrans)":
            translator = GoogleTranslator()
            translated = translator.translate(text, dest=dest, src=src).text
        elif translator_choice == "Microsoft":
            translated = MicrosoftTranslator(source=src, target=dest).translate(text)
        elif translator_choice == "Pons":
            translated = PonsTranslator(source=src, target=dest).translate(text)
        elif translator_choice == "Linguee":
            translated = LingueeTranslator(source=src, target=dest).translate(text)
        elif translator_choice == "MyMemory":
            translated = MyMemoryTranslator(source=src, target=dest).translate(text)
        else:
            raise ValueError("Invalid translator choice")

        return translated

    except Exception as e:
        messagebox.showerror("Translation Error", str(e))
        return ""


def data():
    """ Handles the translation button click event. Retrieves input, performs translation, and updates the UI. """
    s = combo_txt.get()
    d = dest_combo.get()
    masg = sor_txt.get(1.0, END).strip()
    translator_choice = translator_var.get()
    if not masg:
        messagebox.showwarning("Empty Input", "Please enter text to translate.")
        return
    textget = change(text=masg, src=s, dest=d, translator_choice=translator_choice)
    resul_txt.delete(1.0, END)
    resul_txt.insert(END, textget)
    translation_history.append({
        "source": s,
        "destination": d,
        "original": masg,
        "translated": textget,
        "translator": translator_choice
    })
    if history_listbox:
        update_history_listbox()

def swap_languages():
    """ Swaps the source and destination languages and their corresponding texts. """
    src = combo_txt.get()
    dest = dest_combo.get()
    combo_txt.set(dest)
    dest_combo.set(src)
    src_text = sor_txt.get(1.0, END).strip()
    dest_text = resul_txt.get(1.0, END).strip()
    sor_txt.delete(1.0, END)
    resul_txt.delete(1.0, END)
    sor_txt.insert(END, dest_text)
    resul_txt.insert(END, src_text)

def clear_text():
    """ Clears the text from both the source and result Text widgets. """
    sor_txt.delete(1.0, END)
    resul_txt.delete(1.0, END)

def copy_text():
    """ Copies the translated text to the clipboard. """
    text_to_copy = resul_txt.get(1.0, END).strip()
    pyperclip.copy(text_to_copy)
    messagebox.showinfo("Copied", "Translated text copied to clipboard!")

def update_history_listbox():
    """ Updates the Listbox that displays the translation history. """
    if history_listbox and history_listbox.winfo_exists():
        history_listbox.delete(0, END)
        for i, item in enumerate(reversed(translation_history), 1):
            history_listbox.insert(END, f"{i}. {item['original'][:30]}... -> {item['translated'][:30]}...")

def show_history():
    """ Displays the translation history in a new window. """
    global history_listbox
    history_window = Toplevel(box)
    history_window.title("Translation History")
    history_window.geometry("600x400")
    history_frame = Frame(history_window)
    history_frame.pack(fill=BOTH, expand=True)
    history_listbox = Listbox(history_frame, font=("Helvetica", 12))
    history_listbox.pack(side=LEFT, fill=BOTH, expand=True)
    scrollbar = Scrollbar(history_frame, orient=VERTICAL)
    scrollbar.config(command=history_listbox.yview)
    scrollbar.pack(side=RIGHT, fill=Y)
    history_listbox.config(yscrollcommand=scrollbar.set)
    update_history_listbox()

def load_selected_translation():
    """ Loads a selected translation from history into the main window. """
    selection = history_listbox.curselection()
    if selection:
        index = selection[0]
        item = translation_history[-(index + 1)]
        sor_txt.delete(1.0, END)
        sor_txt.insert(END, item["original"])
        resul_txt.delete(1.0, END)
        resul_txt.insert(END, item["translated"])
        combo_txt.set(item["source"])
        dest_combo.set(item["destination"])
        translator_var.set(item["translator"])
        history_window.destroy()

def speech_to_text():
    """ Converts speech to text using Google's speech recognition API. """
    r = sr.Recognizer()
    with sr.Microphone() as source:
        messagebox.showinfo("Speech Recognition", "Speak now...")
        audio = r.listen(source)
    try:
        text = r.recognize_google(audio)
        sor_txt.delete(1.0, END)
        sor_txt.insert(END, text)
    except sr.UnknownValueError:
        messagebox.showerror("Speech Recognition", "Could not understand audio")
    except sr.RequestError as e:
        messagebox.showerror("Speech Recognition", f"Could not request results; {e}")

def text_to_speech():
    """ Converts the translated text to speech. """
    text = resul_txt.get(1.0, END).strip()
    if text:
        tts = gTTS(text=text, lang=dest_combo.get())
        tts.save("translated.mp3")
        os.system("start translated.mp3")  # Adjust for OS
    else:
        messagebox.showwarning("Empty Text", "No text to convert to speech.")

def detect_language():
    """ Detects the language of the input text. """
    text = sor_txt.get(1.0, END).strip()
    if text:
        lang_code = detect(text)
        combo_txt.set(lang_code)  # Set the detected language as the source language
    else:
        messagebox.showwarning("Empty Text", "Please enter text to detect language.")

def import_file():
    """ Imports text from a selected file and inserts it into the source text widget. """
    file_path = filedialog.askopenfilename(
        filetypes=[("Text files", "*.txt"), ("Word files", "*.docx"), ("PDF files", "*.pdf")]
    )
    if file_path:
        ext = file_path.split('.')[-1]
        if ext == "txt":
            with open(file_path, 'r', encoding='utf-8') as file:
                sor_txt.delete(1.0, END)
                sor_txt.insert(END, file.read())
        elif ext == "docx":
            doc = Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            sor_txt.delete(1.0, END)
            sor_txt.insert(END, '\n'.join(full_text))
        elif ext == "pdf":
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfFileReader(file)
                full_text = []
                for page_num in range(reader.numPages):
                    page = reader.getPage(page_num)
                    full_text.append(page.extract_text())
                sor_txt.delete(1.0, END)
                sor_txt.insert(END, '\n'.join(full_text))

def save_translation():
    """ Saves the translated text to a selected file. """
    file_path = filedialog.asksaveasfilename(
        defaultextension=".txt",
        filetypes=[("Text files", "*.txt"), ("Word files", "*.docx")]
    )
    if file_path:
        ext = file_path.split('.')[-1]
        if ext == "txt":
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(resul_txt.get(1.0, END).strip())
        elif ext == "docx":
            doc = Document()
            doc.add_paragraph(resul_txt.get(1.0, END).strip())
            doc.save(file_path)

def save_profile():
    """ Saves the user profile including custom languages and translation history. """
    profile = {
        "custom_languages": custom_languages,
        "translation_history": translation_history
    }
    file_path = filedialog.asksaveasfilename(
        defaultextension=".json",
        filetypes=[("JSON files", "*.json")]
    )
    if file_path:
        with open(file_path, 'w') as file:
            json.dump(profile, file, indent=4)
        messagebox.showinfo("Profile Saved", "User profile has been saved!")

def load_profile():
    """ Loads the user profile including custom languages and translation history. """
    global custom_languages, translation_history
    file_path = filedialog.askopenfilename(
        filetypes=[("JSON files", "*.json")]
    )
    if file_path:
        with open(file_path, 'r') as file:
            profile = json.load(file)
            custom_languages = profile.get("custom_languages", {})
            translation_history = profile.get("translation_history", [])
            update_language_options()
            update_history_listbox()
            messagebox.showinfo("Profile Loaded", "User profile has been loaded!")

def add_custom_language():
    """ Adds a custom language to the language options. """
    def save_language():
        lang_code = lang_code_entry.get().strip()
        lang_name = lang_name_entry.get().strip()
        if lang_code and lang_name:
            custom_languages[lang_code] = lang_name
            update_language_options()
            add_language_window.destroy()
        else:
            messagebox.showwarning("Input Error", "Please provide both language code and name.")

    add_language_window = Toplevel(box)
    add_language_window.title("Add Custom Language")
    Label(add_language_window, text="Language Code:").pack(pady=5)
    lang_code_entry = Entry(add_language_window)
    lang_code_entry.pack(pady=5)
    Label(add_language_window, text="Language Name:").pack(pady=5)
    lang_name_entry = Entry(add_language_window)
    lang_name_entry.pack(pady=5)
    Button(add_language_window, text="Add Language", command=save_language).pack(pady=10)

def update_language_options():
    """ Updates the language options in the dropdown menus with custom languages. """
    src_languages = list(LANGUAGES.keys()) + list(custom_languages.keys())
    dest_languages = list(LANGUAGES.keys()) + list(custom_languages.keys())
    combo_txt.set('english')
    dest_combo.set('marathi')
    menu_src = src_lang_menu['menu']
    menu_dest = dest_lang_menu['menu']
    menu_src.delete(0, END)
    menu_dest.delete(0, END)
    for lang in src_languages:
        menu_src.add_command(label=LANGUAGES.get(lang, lang), command=lambda l=lang: combo_txt.set(l))
    for lang in dest_languages:
        menu_dest.add_command(label=LANGUAGES.get(lang, lang), command=lambda l=lang: dest_combo.set(l))

# Tkinter GUI setup
box = Tk()
box.geometry('500x900')
box.title('Multi-Translator')
box.config(bg='black')

lab_txt = Label(box, text='Multi-Translator', font=('Helvetica', 22, 'bold'), bg='black', fg='white')
lab_txt.pack(pady=20)

frame = Frame(box)
frame.pack(padx=10, pady=10, fill=BOTH, expand=True)

sor_txt = Text(frame, wrap='word', font=('Helvetica', 12), height=10, bg='lightyellow')
sor_txt.pack(side=LEFT, fill=BOTH, expand=True)

resul_txt = Text(frame, wrap='word', font=('Helvetica', 12), height=10, bg='lightblue')
resul_txt.pack(side=RIGHT, fill=BOTH, expand=True)

button_frame = Frame(box)
button_frame.pack(padx=10, pady=10)

translate_button = Button(button_frame, text='Translate', relief=RAISED, command=data, bg='blue', fg='white')
translate_button.pack(side=LEFT, padx=5)

swap_button = Button(button_frame, text='Swap', relief=RAISED, command=swap_languages, bg='green', fg='white')
swap_button.pack(side=LEFT, padx=5)

clear_button = Button(button_frame, text='Clear', relief=RAISED, command=clear_text, bg='red', fg='white')
clear_button.pack(side=LEFT, padx=5)

copy_button = Button(button_frame, text='Copy', relief=RAISED, command=copy_text, bg='purple', fg='white')
copy_button.pack(side=LEFT, padx=5)

history_button = Button(button_frame, text='History', relief=RAISED, command=show_history, bg='orange', fg='white')
history_button.pack(side=LEFT, padx=5)

import_button = Button(button_frame, text='Import File', relief=RAISED, command=import_file, bg='teal', fg='white')
import_button.pack(side=LEFT, padx=5)

save_button = Button(button_frame, text='Save Translation', relief=RAISED, command=save_translation, bg='magenta', fg='white')
save_button.pack(side=LEFT, padx=5)

speech_button = Button(button_frame, text='Speech to Text', relief=RAISED, command=speech_to_text, bg='cyan', fg='black')
speech_button.pack(side=LEFT, padx=5)

text_to_speech_button = Button(button_frame, text='Text to Speech', relief=RAISED, command=text_to_speech, bg='lightgreen', fg='black')
text_to_speech_button.pack(side=LEFT, padx=5)

detect_language_button = Button(button_frame, text='Detect Language', relief=RAISED, command=detect_language, bg='orange', fg='black')
detect_language_button.pack(side=LEFT, padx=5)

profile_frame = Frame(box)
profile_frame.pack(padx=10, pady=10)

save_profile_button = Button(profile_frame, text='Save Profile', relief=RAISED, command=save_profile, bg='grey', fg='white')
save_profile_button.pack(side=LEFT, padx=5)

load_profile_button = Button(profile_frame, text='Load Profile', relief=RAISED, command=load_profile, bg='brown', fg='white')
load_profile_button.pack(side=LEFT, padx=5)

add_language_button = Button(profile_frame, text='Add Custom Language', relief=RAISED, command=add_custom_language, bg='lime', fg='black')
add_language_button.pack(side=LEFT, padx=5)

lang_frame = Frame(box)
lang_frame.pack(padx=10, pady=10)

combo_txt = StringVar()
combo_txt.set('english')
src_lang_menu = OptionMenu(lang_frame, combo_txt, *LANGUAGES.keys())
src_lang_menu.pack(side=LEFT, padx=5)

dest_combo = StringVar()
dest_combo.set('marathi')
dest_lang_menu = OptionMenu(lang_frame, dest_combo, *LANGUAGES.keys())
dest_lang_menu.pack(side=LEFT, padx=5)

translator_var = StringVar()
translator_var.set("Google (deep-translator)")  # Set Google (deep-translator) as the default
translator_menu = OptionMenu(box, translator_var, "Google (googletrans)", "Google (deep-translator)", "Microsoft", "Pons", "Linguee", "MyMemory")
translator_menu.pack(padx=10, pady=10)


update_language_options()
box.mainloop()